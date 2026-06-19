#!/usr/bin/env python3
"""Build an editable Scientific Reports manuscript DOCX from Markdown.

Design preset: narrative_proposal, with a named journal-manuscript override.
Project font override: English/numbers use Times New Roman; East Asian text
uses Songti through the Word eastAsia font slot.
"""

from __future__ import annotations

import argparse
import csv
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ASCII_FONT = "Times New Roman"
EA_FONT = "\u5b8b\u4f53"
MONO_FONT = "Courier New"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(90, 90, 90)


def set_run_font(run, size=None, bold=None, italic=None, color=None, mono=False):
    name = MONO_FONT if mono else ASCII_FONT
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), EA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def configure_style(style, size=11, bold=False, italic=False, color=BLACK):
    style.font.name = ASCII_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), EA_FONT)


def set_para_spacing(paragraph, before=0, after=6, line=1.25, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def set_cell_text_font(cell, size=10):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=size)


def add_page_field(paragraph):
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_inline_markdown(paragraph, text, default_size=11):
    i = 0
    while i < len(text):
        candidates = []
        for marker in ("**", "`", "*"):
            idx = text.find(marker, i)
            if idx != -1:
                candidates.append((idx, marker))
        if not candidates:
            run = paragraph.add_run(text[i:])
            set_run_font(run, size=default_size)
            break
        idx, marker = min(candidates, key=lambda x: x[0])
        if idx > i:
            run = paragraph.add_run(text[i:idx])
            set_run_font(run, size=default_size)
        if marker == "**":
            end = text.find("**", idx + 2)
            if end == -1:
                run = paragraph.add_run(text[idx:])
                set_run_font(run, size=default_size)
                break
            run = paragraph.add_run(text[idx + 2 : end])
            set_run_font(run, size=default_size, bold=True)
            i = end + 2
        elif marker == "`":
            end = text.find("`", idx + 1)
            if end == -1:
                run = paragraph.add_run(text[idx:])
                set_run_font(run, size=default_size)
                break
            run = paragraph.add_run(text[idx + 1 : end])
            set_run_font(run, size=default_size, mono=True)
            i = end + 1
        else:
            end = text.find("*", idx + 1)
            if end == -1:
                run = paragraph.add_run(text[idx:])
                set_run_font(run, size=default_size)
                break
            run = paragraph.add_run(text[idx + 1 : end])
            set_run_font(run, size=default_size, italic=True)
            i = end + 1


def add_title(doc, title):
    p = doc.add_paragraph()
    p.style = "Manuscript Title"
    set_para_spacing(p, before=0, after=10, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(title)
    set_run_font(run, size=16, bold=True)

    p2 = doc.add_paragraph()
    set_para_spacing(p2, before=0, after=14, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p2.add_run("Scientific Reports article draft | public-data bioinformatics | author metadata supplied")
    set_run_font(run, size=10, color=MUTED)


def add_heading(doc, text, level):
    style = "Manuscript Heading 1" if level == 1 else "Manuscript Heading 2"
    p = doc.add_paragraph()
    p.style = style
    if level == 1:
        set_para_spacing(p, before=14, after=6, line=1.15)
        size = 13
    else:
        set_para_spacing(p, before=10, after=4, line=1.15)
        size = 11.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)


def add_body_paragraph(doc, text, in_references=False):
    p = doc.add_paragraph()
    p.style = "Reference Paragraph" if in_references else "Normal"
    if in_references:
        set_para_spacing(p, before=0, after=5, line=1.15)
    else:
        set_para_spacing(p, before=0, after=7, line=1.28, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_inline_markdown(p, text, default_size=10.8 if in_references else 11)


def read_markdown_blocks(md_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    blocks = []
    buf = []
    standalone_patterns = (
        re.compile(r"^\d+\.\s+"),
        re.compile(r"^\*\*(Main Figure|Figure|Extended Data Figure)\s+\d+\."),
    )

    def flush():
        nonlocal buf
        if buf:
            blocks.append(("p", " ".join(x.strip() for x in buf).strip()))
            buf = []

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush()
            continue
        if line.startswith("# "):
            flush()
            blocks.append(("title", line[2:].strip()))
        elif line.startswith("## "):
            flush()
            blocks.append(("h1", line[3:].strip()))
        elif line.startswith("### "):
            flush()
            blocks.append(("h2", line[4:].strip()))
        elif any(pattern.match(line.strip()) for pattern in standalone_patterns):
            flush()
            blocks.append(("p", line.strip()))
        else:
            buf.append(line)
    flush()
    return blocks


def configure_document(doc: Document):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    configure_style(normal, size=11, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.28

    styles = doc.styles
    title_style = styles.add_style("Manuscript Title", 1)
    configure_style(title_style, size=16, bold=True, color=BLACK)

    h1 = styles.add_style("Manuscript Heading 1", 1)
    configure_style(h1, size=13, bold=True, color=BLACK)
    h1.paragraph_format.keep_with_next = True

    h2 = styles.add_style("Manuscript Heading 2", 1)
    configure_style(h2, size=11.5, bold=True, color=BLACK)
    h2.paragraph_format.keep_with_next = True

    ref_style = styles.add_style("Reference Paragraph", 1)
    configure_style(ref_style, size=10.8, color=BLACK)
    ref_style.paragraph_format.left_indent = Inches(0.28)
    ref_style.paragraph_format.first_line_indent = Inches(-0.28)
    ref_style.paragraph_format.space_after = Pt(5)
    ref_style.paragraph_format.line_spacing = 1.15

    header = section.header
    if header.paragraphs:
        hp = header.paragraphs[0]
    else:
        hp = header.add_paragraph()
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(hp, before=0, after=0, line=1.0)
    hr = hp.add_run("TMEM158 Ca2/UPR-CAF stress ecology | Scientific Reports draft")
    set_run_font(hr, size=8.5, color=MUTED)

    footer = section.footer
    if footer.paragraphs:
        fp = footer.paragraphs[0]
    else:
        fp = footer.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(fp, before=0, after=0, line=1.0)
    add_page_field(fp)
    for run in fp.runs:
        set_run_font(run, size=8.5, color=MUTED)


def build_docx(md_path: Path, out_path: Path, qc_path: Path):
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "TMEM158-associated Ca2/UPR-CAF stress ecology state in ESCC"
    doc.core_properties.subject = "Scientific Reports article draft"
    doc.core_properties.author = "Yang Haoshui; Ma Yuqing"
    doc.core_properties.comments = "Generated from manuscript_scientific_reports.md; author metadata supplied; publisher preview remains human-gated."

    in_references = False
    title_seen = False
    paragraph_count = 0
    heading_count = 0
    reference_count = 0

    for kind, text in read_markdown_blocks(md_path):
        if kind == "title":
            add_title(doc, text)
            title_seen = True
        elif kind == "h1":
            if text.strip().lower() == "references":
                in_references = True
            add_heading(doc, text, level=1)
            heading_count += 1
        elif kind == "h2":
            add_heading(doc, text, level=2)
            heading_count += 1
        else:
            if in_references and re.match(r"^\d+\.\s+", text):
                reference_count += 1
            add_body_paragraph(doc, text, in_references=in_references)
            paragraph_count += 1

    if not title_seen:
        raise RuntimeError("No markdown title found.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    qc_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    rows = [
        ("source_markdown", str(md_path)),
        ("output_docx", str(out_path)),
        ("design_preset", "narrative_proposal"),
        ("named_override", "journal_manuscript_times_new_roman_songti"),
        ("page_size", "US Letter portrait"),
        ("margins", "1 inch"),
        ("ascii_font", ASCII_FONT),
        ("east_asia_font", EA_FONT),
        ("body_size_pt", "11"),
        ("line_spacing", "1.28 body; 1.15 references"),
        ("title_seen", str(title_seen)),
        ("heading_count", str(heading_count)),
        ("paragraph_count", str(paragraph_count)),
        ("reference_count", str(reference_count)),
        ("generated_at", datetime.now().isoformat(timespec="seconds")),
    ]
    with qc_path.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(["item", "value"])
        writer.writerows(rows)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--input",
        default="TMEM158_CaUPR_ESCC/07_manuscript/manuscript_scientific_reports.md",
    )
    parser.add_argument(
        "--output",
        default="TMEM158_CaUPR_ESCC/07_manuscript/manuscript_scientific_reports.docx",
    )
    parser.add_argument(
        "--qc",
        default="TMEM158_CaUPR_ESCC/04_results/qc/scientific_reports_docx_build_qc.csv",
    )
    args = parser.parse_args()
    build_docx(Path(args.input), Path(args.output), Path(args.qc))
    print(f"Wrote {args.output}")
    print(f"Wrote {args.qc}")


if __name__ == "__main__":
    main()
